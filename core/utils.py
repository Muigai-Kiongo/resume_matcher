"""
Resume parsing utilities.

Improvements:
- More robust handling when spaCy or pdfplumber are not installed.
- Accepts either a file path or a file-like object for text extraction.
- Safer skill extraction with normalization and order-preserving deduplication.
- Experience/education extraction uses spaCy where available, falls back to regex/heuristic parsing.
- calculate_match_score accepts lists of strings, QuerySets of Skill objects, or comma-separated strings.
"""
from typing import List, Iterable, Optional, Union
import re
import warnings

# Optional imports: handle gracefully so the whole app doesn't crash if a package is missing.
try:
    import pdfplumber
except Exception:  # pragma: no cover - defensive
    pdfplumber = None
try:
    import docx
except Exception:  # pragma: no cover - defensive
    docx = None

# spaCy may not be installed or the model may not be available in all environments.
nlp = None
try:
    import spacy

    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        # If loading the model fails, fallback to None and continue with heuristics.
        nlp = None
        warnings.warn("spaCy model en_core_web_sm not available. Falling back to heuristics.")
except Exception:
    nlp = None
    warnings.warn("spaCy not installed. Falling back to heuristics for NLP tasks.")


def _read_file_text(file_path_or_file: Union[str, "os.PathLike", object], ext: Optional[str] = None) -> str:
    """
    Low-level helper to obtain text from either a file path (string) or a file-like object
    (e.g., Django InMemoryUploadedFile). `ext` may be provided to skip extension detection.
    """
    # Try to detect extension if not provided
    if ext is None:
        try:
            name = getattr(file_path_or_file, "name", None)
            if name:
                ext = name.split(".")[-1].lower()
        except Exception:
            ext = None

    if isinstance(file_path_or_file, (str,)):
        path = file_path_or_file
        if ext == "pdf" or (ext is None and path.lower().endswith(".pdf")):
            if not pdfplumber:
                warnings.warn("pdfplumber not installed; cannot extract text from PDF.")
                return ""
            text = ""
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        elif ext in ("doc", "docx") or (ext is None and path.lower().endswith((".doc", ".docx"))):
            if not docx:
                warnings.warn("python-docx not installed; cannot extract text from DOCX.")
                return ""
            document = docx.Document(path)
            return "\n".join([p.text for p in document.paragraphs])
        else:
            # Unknown file type - try to open as text
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                    return fh.read()
            except Exception:
                return ""
    else:
        # Assume file-like object (UploadedFile). Read bytes and branch by ext or attempt pdf/docx detection.
        uploaded = file_path_or_file
        name = getattr(uploaded, "name", "")
        guessed_ext = (ext or (name.split(".")[-1].lower() if name and "." in name else None))
        uploaded.seek(0)
        content = uploaded.read()
        # If bytes, work accordingly
        if isinstance(content, bytes):
            if guessed_ext == "pdf" or (guessed_ext is None and content[:4] == b"%PDF"):
                if not pdfplumber:
                    warnings.warn("pdfplumber not installed; cannot extract text from PDF.")
                    return ""
                import io

                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
            elif guessed_ext in ("doc", "docx") or (guessed_ext is None and content[:2] == b'PK'):
                # DOCX files are ZIP/PK archives
                if not docx:
                    warnings.warn("python-docx not installed; cannot extract text from DOCX.")
                    return ""
                import io

                document = docx.Document(io.BytesIO(content))
                return "\n".join([p.text for p in document.paragraphs])
            else:
                # Try to decode as text
                try:
                    return content.decode("utf-8", errors="ignore")
                except Exception:
                    return ""
        else:
            # content probably already str
            return str(content)


def extract_text_from_resume(file_path_or_file: Union[str, object], ext: Optional[str] = None) -> str:
    """
    Public helper to extract text from resume file.
    Accepts:
      - file path (string)
      - file-like object (UploadedFile)
    ext: optional extension hint (e.g., 'pdf', 'docx')
    """
    try:
        return _read_file_text(file_path_or_file, ext) or ""
    except Exception as e:
        warnings.warn(f"Error extracting text from resume: {e}")
        return ""


def _normalize_items(items: Optional[Iterable]) -> List[str]:
    """Normalize an iterable of items to a list of unique, stripped strings (preserve order)."""
    if not items:
        return []
    out = []
    seen = set()
    for it in items:
        if it is None:
            continue
        s = str(it).strip()
        if not s:
            continue
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s)
    return out


def extract_skills(text: str, skills_list: Optional[List[str]] = None) -> List[str]:
    """
    Extracts known skills from text. If skills_list is provided, it will search for these
    items in the text (case-insensitive). Otherwise uses a conservative default set.

    Returns a list of distinct skill strings in the order found in skills_list (or discovered).
    """
    if not text:
        return []
    if skills_list:
        candidates = skills_list
    else:
        candidates = [
            "Python", "Java", "Excel", "Machine Learning", "Django", "React", "SQL",
            "Project Management", "Communication", "Leadership", "AWS", "Docker", "Kubernetes",
            "REST", "GraphQL", "TypeScript", "JavaScript", "HTML", "CSS"
        ]

    text_lower = text.lower()
    found = []
    for skill in candidates:
        if skill and skill.lower() in text_lower:
            found.append(skill)
    # As a fallback, try to find capitalized words sequences that look like skills (e.g., "TensorFlow", "PyTorch")
    if not found:
        # regex to capture words with camelcase or mixed caps, or known patterns
        pattern = r"\b[A-Z][a-zA-Z0-9\+\#\.\-]{2,}\b"
        matches = re.findall(pattern, text)
        found += _normalize_items(matches)
    # Deduplicate preserving order and normalize whitespace
    return _normalize_items(found)


def extract_experience(text: str, limit: int = 5) -> List[str]:
    """
    Extracts likely experience paragraphs/sentences. If spaCy is available, use sentence segmentation
    and heuristics. Otherwise uses simple regex that looks for lines containing 'experience', years, or 'worked at'.
    Returns up to `limit` entries (most relevant first).
    """
    if not text:
        return []
    results = []
    if nlp:
        doc = nlp(text)
        for sent in doc.sents:
            s = sent.text.strip()
            lower = s.lower()
            # heuristics
            if ("experience" in lower or "worked at" in lower or "responsible for" in lower or re.search(r"\b\d{4}\b", s)):
                results.append(s)
                if len(results) >= limit:
                    break
    else:
        # fallback: split by newlines and look for lines with date ranges or keywords
        lines = re.split(r"\n{1,}", text)
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if ("experience" in s.lower() or "worked at" in s.lower() or re.search(r"\b(19|20)\d{2}\b", s)):
                results.append(s)
                if len(results) >= limit:
                    break
    return results


def extract_education(text: str) -> Optional[str]:
    """
    Attempts to find the most relevant education line/paragraph.
    Uses spaCy NER (if available) to find ORG and dates, otherwise uses regex for degree keywords.
    """
    if not text:
        return None
    if nlp:
        doc = nlp(text)
        # Search for sentences with university/college/degree keywords
        for sent in doc.sents:
            s = sent.text.strip()
            if re.search(r"\b(university|college|bachelor|master|degree|phd|mba)\b", s, re.I):
                return s
        # As fallback, search entities labelled ORG nearby Degree words
        for ent in doc.ents:
            if ent.label_ in ("ORG", "PERSON") and re.search(r"\b(university|college)\b", ent.text, re.I):
                return ent.text
    else:
        # simple heuristic: look line by line for degree/university mentions
        for line in text.splitlines():
            if re.search(r"\b(university|college|bachelor|master|degree|phd|mba)\b", line, re.I):
                return line.strip()
    return None


def calculate_match_score(resume_skills: Union[Iterable, str, None],
                          job_requirements: Union[Iterable, str, None]) -> float:
    """
    Calculate match score between resume_skills and job_requirements.
    Accepts:
      - iterable of strings (skill names)
      - queryset/list of Skill model instances (they will be converted to .name)
      - comma-separated string
    Returns a float between 0 and 1 (rounded to 2 decimals). If job_requirements is empty, returns 0.0.
    """
    def to_names(x):
        if x is None:
            return []
        if isinstance(x, str):
            # split comma separated
            parts = [p.strip() for p in re.split(r",|;", x) if p.strip()]
            return parts
        try:
            # Try iterating and extracting .name attr when present
            out = []
            for it in x:
                if hasattr(it, "name"):
                    out.append(str(it.name))
                else:
                    out.append(str(it))
            return out
        except TypeError:
            return []

    resume_list = _normalize_items(to_names(resume_skills))
    job_list = _normalize_items(to_names(job_requirements))

    if not job_list:
        return 0.0
    resume_set = set([s.lower() for s in resume_list])
    job_set = set([s.lower() for s in job_list])
    overlap = resume_set.intersection(job_set)
    score = len(overlap) / len(job_set)
    return round(score, 2)